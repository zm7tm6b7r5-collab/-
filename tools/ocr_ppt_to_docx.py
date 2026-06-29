r"""
品牌管理课程 PPT 图片 OCR -> Word 文档
用法: python ocr_ppt_to_docx.py
输入: 85张PPT拍照图片 (JPG)
输出: 品牌管理课程笔记.docx
"""

import subprocess
import os
import sys
import re
from pathlib import Path
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ===== 配置 =====
IMG_DIR = Path(r"C:\Users\AUSU\Desktop\品牌")
OUTPUT_PATH = Path(r"C:\Users\AUSU\Desktop\品牌管理课程笔记.docx")
TESSERACT_EXE = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSDATA_DIR = os.path.expandvars(r"%USERPROFILE%\tesseract-tessdata")
LANG = "chi_sim"  # 中文简体

# ===== 步骤 1: 获取所有图片并按文件名排序 =====
def get_images():
    """获取所有 JPG 图片"""
    if not IMG_DIR.exists():
        print(f"❌ 图片目录不存在: {IMG_DIR}")
        sys.exit(1)
    images = sorted(IMG_DIR.glob("*.jpg"))
    print(f"📸 找到 {len(images)} 张图片")
    return images


# ===== 步骤 2: 预处理图片（增强对比度，提高 OCR 准确率）=====
def preprocess_image(img_path: Path) -> Path:
    """
    对图片进行预处理以提高OCR准确率:
    - 转换为灰度
    - 增强对比度
    - 二值化
    返回预处理后的临时图片路径
    """
    img = Image.open(img_path)
    # 转换为灰度
    gray = img.convert("L")
    # 简单对比度增强：使用point操作
    # 拉伸直方图
    enhanced = gray.point(lambda x: 0 if x < 60 else (255 if x > 200 else int((x - 60) * 255 / 140)))
    # 保存临时文件
    temp_path = img_path.parent / f"_temp_preprocessed_{img_path.stem}.png"
    enhanced.save(temp_path, "PNG")
    return temp_path


# ===== 步骤 3: OCR 单张图片 =====
def ocr_image(img_path: Path, use_preprocess: bool = True) -> str:
    """
    对单张图片进行OCR识别
    返回识别出的文本
    """
    process_path = img_path
    temp_path = None

    if use_preprocess:
        try:
            temp_path = preprocess_image(img_path)
            process_path = temp_path
        except Exception as e:
            print(f"  ⚠️ 预处理失败，使用原图: {e}")
            process_path = img_path

    env = os.environ.copy()
    env["TESSDATA_PREFIX"] = TESSDATA_DIR

    try:
        result = subprocess.run(
            [TESSERACT_EXE, str(process_path), "stdout", "-l", LANG, "--psm", "6"],
            capture_output=True,
            text=True,
            env=env,
            timeout=60,
            encoding="utf-8",
        )
        # 清理临时文件
        if temp_path and temp_path.exists():
            temp_path.unlink()

        if result.returncode != 0:
            stderr = result.stderr.strip()
            # 忽略 resolution 估算警告
            if stderr and "Estimating resolution" not in stderr:
                print(f"  ⚠️ OCR 警告: {stderr[:100]}")
        return result.stdout.strip() if result.stdout else ""
    except subprocess.TimeoutExpired:
        print(f"  ❌ OCR 超时")
        if temp_path and temp_path.exists():
            temp_path.unlink()
        return ""
    except Exception as e:
        print(f"  ❌ OCR 异常: {e}")
        if temp_path and temp_path.exists():
            temp_path.unlink()
        return ""


# ===== 步骤 4: 清理 OCR 文本 =====
def clean_text(text: str) -> str:
    """清理OCR产生的噪音和格式问题"""
    # 去除多余空行（保留最多1个连续空行）
    text = re.sub(r"\n{3,}", "\n\n", text)
    # 去除每行首尾多余空格
    lines = [line.strip() for line in text.split("\n")]
    # 去除完全空白的行首尾
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)


# ===== 步骤 5: 分析文本判断是否为 PPT 封面页 =====
def is_title_slide(text: str) -> bool:
    """根据文本特征判断是否为PPT封面/标题页"""
    lines = [l for l in text.split("\n") if l.strip()]
    if not lines:
        return False
    # 封面通常文字少、字号大（OCR后可能表现为短行）
    total_chars = sum(len(l) for l in lines)
    if total_chars < 50 and len(lines) <= 5:
        return True
    # 包含"第X章" "第X讲" 等特征
    for l in lines[:3]:
        if re.search(r"第[一二三四五六七八九十\d]+[章节讲]", l):
            return True
    return False


# ===== 步骤 6: 生成 Word 文档 =====
def create_docx(all_texts: list, doc_path: Path):
    """将所有OCR文本整理为Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 标题
    title = doc.add_heading("品牌管理课程笔记", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"共 {len(all_texts)} 页 | OCR 自动识别 | 整理日期：2026-06-26")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 按页输出
    page_num = 0
    for i, text in enumerate(all_texts):
        if not text or not text.strip():
            continue
        page_num += 1

        # 清理文本
        cleaned = clean_text(text)

        # 封面页用大标题
        if is_title_slide(cleaned):
            lines = cleaned.split("\n")
            for line in lines:
                if line.strip():
                    heading = doc.add_heading(line.strip(), level=1)
        else:
            # 普通页：加页码标记
            para_page = doc.add_paragraph()
            run_page = para_page.add_run(f"—— 第 {page_num} 页 ——")
            run_page.font.size = Pt(9)
            run_page.font.color.rgb = RGBColor(150, 150, 150)

            # 分行添加内容
            lines = cleaned.split("\n")
            current_para = None
            for line in lines:
                stripped = line.strip()
                if not stripped:
                    # 空行表示段落分隔
                    current_para = None
                    continue

                # 判断是否为新段落（短行可能是标题）
                if len(stripped) <= 30 and (
                    re.search(r"^[一二三四五六七八九十\d]+[、.)\s]", stripped)
                    or re.search(r"^(第[一二三四五六七八九十\d]+[章节])", stripped)
                    or re.search(r"^[A-Z][A-Za-z\s]{2,20}$", stripped)
                ):
                    # 这是一个小标题
                    doc.add_heading(stripped, level=3)
                    current_para = None
                else:
                    # 正文内容
                    if current_para is None:
                        current_para = doc.add_paragraph()
                    else:
                        current_para.add_run("\n")
                    run = current_para.add_run(stripped)

    # 保存
    try:
        doc.save(str(doc_path))
        print(f"\n✅ 文档已保存: {doc_path}")
    except Exception as e:
        print(f"❌ 保存失败: {e}")
        # 尝试保存到其他位置
        alt_path = Path.home() / "Desktop" / "品牌管理课程笔记.docx"
        doc.save(str(alt_path))
        print(f"✅ 已保存到备用位置: {alt_path}")


# ===== 主流程 =====
def main():
    print("=" * 60)
    print("  品牌管理 PPT → Word 文档 OCR 转换工具")
    print("=" * 60)

    # 1. 获取图片
    images = get_images()
    if not images:
        print("❌ 没有找到图片文件")
        sys.exit(1)

    # 2. 逐张 OCR
    all_texts = []
    for idx, img_path in enumerate(images, 1):
        print(f"\n[{idx}/{len(images)}] OCR 处理: {img_path.name}")
        text = ocr_image(img_path, use_preprocess=True)

        if text:
            preview = text[:100].replace("\n", " | ")
            print(f"  📝 识别 {len(text)} 字符: {preview}...")
        else:
            print(f"  ⚠️ 未识别到文字")

        all_texts.append(text)

    # 3. 统计
    total_chars = sum(len(t) for t in all_texts)
    non_empty = sum(1 for t in all_texts if t.strip())
    print(f"\n📊 统计: {non_empty}/{len(images)} 页识别到文字，共 {total_chars} 字符")

    # 4. 生成 Word 文档
    print("\n📝 正在生成 Word 文档...")
    create_docx(all_texts, OUTPUT_PATH)

    print("\n" + "=" * 60)
    print("  🎉 转换完成！")
    print(f"  输出文件: {OUTPUT_PATH}")
    print("=" * 60)


if __name__ == "__main__":
    main()
