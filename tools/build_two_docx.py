"""
将两个考研英语经济类词汇 Markdown 文件转换为精美 Word 文档
输出到桌面
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from pathlib import Path

# ── 配置 ─────────────────────────────────────────────
DESKTOP = Path.home() / "Desktop"
PROJECT = Path(r"c:\Users\AUSU\Documents\trae_projects\cc")

FILES = [
    PROJECT / "考研英语经济类核心词汇笔记（详细版200词）.md",
    PROJECT / "考研英语经济类核心词汇笔记（续）.md",
]

# 颜色方案
BLUE_DARK   = RGBColor(0x1B, 0x3A, 0x5C)   # 深蓝 — 一级标题
BLUE_MID    = RGBColor(0x2C, 0x5F, 0x8A)   # 中蓝 — 二级标题
BLUE_LIGHT  = RGBColor(0x3A, 0x7C, 0xBF)   # 浅蓝 — 三级标题
GRAY_BG     = "F2F4F7"                       # 灰底 — 表格/提示
ACCENT_GOLD = RGBColor(0xD4, 0x8B, 0x2C)    # 金色 — 强调/icon
RED_WARN    = RGBColor(0xC0, 0x39, 0x2B)    # 红 — 陷阱
GREEN_TIP   = RGBColor(0x27, 0xAE, 0x60)    # 绿 — 提示
TEXT_DARK    = RGBColor(0x2D, 0x2D, 0x2D)    # 正文深灰
TEXT_MEDIUM  = RGBColor(0x55, 0x55, 0x55)    # 次要文字

# ═══════════════════════════════════════════════════════
# 样式系统
# ═══════════════════════════════════════════════════════

def setup_styles(doc: Document):
    """定义全局样式"""
    style = doc.styles['Normal']
    style.font.name = '等线'
    style.font.size = Pt(10.5)
    style.font.color.rgb = TEXT_DARK
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.35
    # 东亚字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # Heading 1
    h1 = doc.styles['Heading 1']
    h1.font.name = '微软雅黑'
    h1.font.size = Pt(22)
    h1.font.bold = True
    h1.font.color.rgb = BLUE_DARK
    h1.paragraph_format.space_before = Pt(28)
    h1.paragraph_format.space_after = Pt(14)
    h1.paragraph_format.keep_with_next = True
    h1.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Heading 2
    h2 = doc.styles['Heading 2']
    h2.font.name = '微软雅黑'
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.font.color.rgb = BLUE_MID
    h2.paragraph_format.space_before = Pt(22)
    h2.paragraph_format.space_after = Pt(10)
    h2.paragraph_format.keep_with_next = True
    h2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Heading 3
    h3 = doc.styles['Heading 3']
    h3.font.name = '微软雅黑'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = BLUE_LIGHT
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True
    h3.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def set_cell_shading(cell, color: str):
    """给单元格加背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc: Document, headers: list, rows: list, col_widths=None):
    """创建美观表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        set_cell_shading(cell, "1B3A5C")

    # 数据行
    for r, row in enumerate(rows):
        bg = "F7F9FC" if r % 2 == 0 else "FFFFFF"
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.color.rgb = TEXT_DARK
            run.font.name = '等线'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
            set_cell_shading(cell, bg)

    # 列宽
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Cm(w)

    doc.add_paragraph()  # 表后空行
    return table


def add_tip_box(doc: Document, text: str, icon: str = "💡"):
    """添加提示框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)

    # 图标
    run_icon = p.add_run(f"{icon}  ")
    run_icon.font.size = Pt(10)
    # 文字
    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = BLUE_MID
    run_text.font.italic = True
    run_text.font.name = '等线'
    run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    # 给段落加左边框效果 — 用缩进 + 灰背景模拟
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F4F8"/>')
    pPr.append(shd)


def add_trap_box(doc: Document, text: str):
    """添加陷阱/警告框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.5)

    run_icon = p.add_run("⚠️  ")
    run_icon.font.size = Pt(10)
    run_text = p.add_run(text)
    run_text.font.size = Pt(9.5)
    run_text.font.color.rgb = RED_WARN
    run_text.font.bold = True
    run_text.font.name = '等线'
    run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="FEF2F2"/>')
    pPr.append(shd)


def add_divider(doc: Document):
    """添加分隔线"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="6" w:space="1" w:color="D0D5DD"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════
# Markdown → Word 转换器
# ═══════════════════════════════════════════════════════

class MD2Docx:
    """逐行解析 Markdown，写入 python-docx Document"""

    def __init__(self, doc: Document):
        self.doc = doc
        self.in_code_block = False
        self.in_table = False
        self.table_rows = []
        self.code_lines = []

    def parse(self, text: str):
        """入口：解析整个文本"""
        lines = text.split('\n')
        i = 0
        while i < len(lines):
            line = lines[i]

            # 代码块
            if line.strip().startswith('```'):
                if self.in_code_block:
                    self._flush_code_block()
                else:
                    self.in_code_block = True
                    self.code_lines = []
                i += 1
                continue

            if self.in_code_block:
                self.code_lines.append(line)
                i += 1
                continue

            # 空行
            if not line.strip():
                i += 1
                continue

            # 表格
            if '|' in line and line.strip().startswith('|'):
                i = self._parse_table(lines, i)
                continue

            # 标题
            h_match = re.match(r'^(#{1,6})\s+(.+)$', line)
            if h_match:
                level = len(h_match.group(1))
                text = h_match.group(2).strip()
                self._add_heading(text, level)
                i += 1
                continue

            # 水平线
            if re.match(r'^[-*_]{3,}\s*$', line.strip()):
                add_divider(self.doc)
                i += 1
                continue

            # 引用块
            if line.strip().startswith('>'):
                i = self._parse_blockquote(lines, i)
                continue

            # 无序列表
            if re.match(r'^(\s*)[-*+]\s+', line):
                i = self._parse_list(lines, i, ordered=False)
                continue

            # 有序列表
            if re.match(r'^(\s*)\d+[.)]\s+', line):
                i = self._parse_list(lines, i, ordered=True)
                continue

            # 普通段落
            i = self._parse_paragraph(lines, i)

        # 收尾
        if self.in_code_block:
            self._flush_code_block()

    # ── 内部方法 ───────────────────────────────────

    def _add_heading(self, text: str, level: int):
        """添加标题"""
        text = self._clean_inline(text)
        level = min(level, 3)  # 最多到 H3
        self.doc.add_heading(text, level=level)

    def _clean_inline(self, text: str) -> str:
        """清理行内标记，保留纯文本"""
        # 粗体 **text** → text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        # 斜体 *text*
        text = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'\1', text)
        # 行内代码 `code`
        text = re.sub(r'`([^`]+)`', r'\1', text)
        # HTML 标签
        text = re.sub(r'<[^>]+>', '', text)
        # 图片
        text = re.sub(r'!\[.*?\]\(.*?\)', '', text)
        # 链接 [text](url)
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        return text.strip()

    def _add_rich_paragraph(self, text: str, style=None):
        """添加支持粗体/斜体/代码的富文本段落"""
        if style:
            p = self.doc.add_paragraph(style=style)
        else:
            p = self.doc.add_paragraph()

        # 简单解析：**bold**, *italic*, `code`
        parts = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
            elif part.strip():
                run = p.add_run(part)
        return p

    def _parse_table(self, lines: list, start: int) -> int:
        """解析 Markdown 表格"""
        rows = []
        i = start
        while i < len(lines) and '|' in lines[i] and lines[i].strip().startswith('|'):
            line = lines[i].strip()
            # 跳过分隔行
            if re.match(r'^[\|\s\-:]+$', line):
                i += 1
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
            i += 1

        if rows:
            headers = rows[0]
            data = rows[1:] if len(rows) > 1 else []
            # 清理行内标记
            headers = [self._clean_inline(h) for h in headers]
            data = [[self._clean_inline(c) for c in r] for r in data]
            add_styled_table(self.doc, headers, data)

        return i

    def _parse_blockquote(self, lines: list, start: int) -> int:
        """解析引用块"""
        i = start
        parts = []
        while i < len(lines) and lines[i].strip().startswith('>'):
            t = lines[i].strip()[1:].strip()
            if t:
                parts.append(t)
            i += 1

        text = ' '.join(parts)
        text = self._clean_inline(text)
        if text:
            # 判断是提示还是警告
            if '⚠️' in text or '陷阱' in text:
                add_trap_box(self.doc, text)
            else:
                add_tip_box(self.doc, text)
        return i

    def _parse_list(self, lines: list, start: int, ordered: bool) -> int:
        """解析列表"""
        i = start
        while i < len(lines):
            line = lines[i]
            if ordered:
                m = re.match(r'^(\s*)\d+[.)]\s+(.+)', line)
            else:
                m = re.match(r'^(\s*)[-*+]\s+(.+)', line)

            if not m:
                break

            indent = len(m.group(1))
            text = m.group(2).strip()

            p = self.doc.add_paragraph(style='List Bullet')
            p.clear()
            self._add_rich_to_paragraph(p, text)
            i += 1

            # 继续行（缩进的内容）
            while i < len(lines) and lines[i].startswith('  ') and not re.match(r'^(\s*)[-*+]\s+', lines[i]) and not re.match(r'^(\s*)\d+[.)]\s+', lines[i]) and not lines[i].strip().startswith('|') and not lines[i].strip().startswith('>') and not lines[i].strip().startswith('#'):
                cont = lines[i].strip()
                if cont:
                    p2 = self.doc.add_paragraph(style='List Bullet')
                    p2.clear()
                    self._add_rich_to_paragraph(p2, cont)
                i += 1

        return i

    def _add_rich_to_paragraph(self, p, text: str):
        """向已存在的段落添加富文本"""
        parts = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
            elif part.strip():
                run = p.add_run(part)

    def _parse_paragraph(self, lines: list, start: int) -> int:
        """解析普通段落"""
        i = start
        parts = []
        while i < len(lines) and lines[i].strip() and \
              not lines[i].strip().startswith('#') and \
              not lines[i].strip().startswith('|') and \
              not lines[i].strip().startswith('>') and \
              not re.match(r'^(\s*)[-*+]\s+', lines[i]) and \
              not re.match(r'^(\s*)\d+[.)]\s+', lines[i]) and \
              not lines[i].strip().startswith('```') and \
              not re.match(r'^[-*_]{3,}\s*$', lines[i].strip()):
            parts.append(lines[i].strip())
            i += 1

        text = ' '.join(parts)
        if text.strip():
            self._add_rich_paragraph(text)
        return i

    def _flush_code_block(self):
        """输出代码块"""
        self.in_code_block = False
        if self.code_lines:
            for cl in self.code_lines:
                p = self.doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(cl)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # 底部空行
            self.doc.add_paragraph()
        self.code_lines = []


# ═══════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════

def add_cover(doc: Document, title: str, subtitle: str):
    """添加精美封面"""
    # 空行推下去
    for _ in range(6):
        doc.add_paragraph()

    # 主标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = '微软雅黑'
    run.font.size = Pt(28)
    run.bold = True
    run.font.color.rgb = BLUE_DARK
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 副标题
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.font.name = '微软雅黑'
    run2.font.size = Pt(14)
    run2.font.color.rgb = BLUE_MID
    run2.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 分隔装饰
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('─' * 40)
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

    # 信息
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run('考研英语 · 经济类核心词汇\n2026 精简版 · 基于 2004-2026 真题高频词')
    run4.font.size = Pt(10.5)
    run4.font.color.rgb = TEXT_MEDIUM
    run4.font.name = '等线'
    run4.element.rPr.rFonts.set(qn('w:eastAsia'), '等线')

    doc.add_page_break()


# ═══════════════════════════════════════════════════════
# 主流程
# ═══════════════════════════════════════════════════════

def convert_file(md_path: Path, output_name: str, title: str, subtitle: str):
    """转换单个 MD 文件为精美 Word 文档"""
    print(f"[处理] {md_path.name}")

    # 读取
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # 创建文档
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # 样式
    setup_styles(doc)

    # 封面
    add_cover(doc, title, subtitle)

    # 转换正文
    converter = MD2Docx(doc)
    converter.parse(content)

    # 保存
    output_path = DESKTOP / output_name
    doc.save(str(output_path))
    print(f"[OK] 已保存: {output_path}")
    return output_path


# ── 执行 ─────────────────────────────────────────────
if __name__ == '__main__':
    print("=" * 50)
    print("  考研英语经济类词汇 → 精美 Word 文档生成器")
    print("=" * 50)

    results = []

    # 文件 1：详细版 200 词
    r1 = convert_file(
        md_path=FILES[0],
        output_name="考研英语经济类核心词汇笔记（详细版1-110词）.docx",
        title="考研英语经济类核心词汇笔记",
        subtitle="详细版 · 第 1–110 词 · 完整词根词缀 + 真题语境 + 阅读陷阱",
    )
    results.append(r1)

    # 文件 2：续（111-200 词）
    r2 = convert_file(
        md_path=FILES[1],
        output_name="考研英语经济类核心词汇笔记（详细版111-200词）.docx",
        title="考研英语经济类核心词汇笔记（续）",
        subtitle="详细版 · 第 111–200 词 · 科技经济 + 消费市场 + 政策制度 + 熟词僻义",
    )
    results.append(r2)

    print("=" * 50)
    print("  [Done] 全部完成！文件位于桌面：")
    for r in results:
        print(f"     >> {r.name}")
    print("=" * 50)
