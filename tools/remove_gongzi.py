# -*- coding: utf-8 -*-
"""
删除 Part2 最终数据库设计中的「工资表（gongzi）」
同时将「五张表」更新为「四张表」
"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from lxml import etree
import shutil
from datetime import datetime

SRC = r'C:\Users\AUSU\Desktop\数据库原理期末大作业(3).docx'

# 备份
backup = SRC.replace('.docx', f'_backup_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
shutil.copy2(SRC, backup)
print(f'✅ 备份已保存: {backup}')

doc = Document(SRC)
body = doc.element.body
nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

children = list(body)

# ==== 第一步：找到需要操作的元素索引 ====
target_indices = {
    'update_176': None,   # "五张表" 段落
    'update_225': None,   # "规范化总结" 段落
    'delete_start': None, # 工资表 heading [204]
    'delete_end': None,   # 空行 [224]
}

for ci, child in enumerate(children):
    tag = child.tag.split('}')[-1]
    if tag != 'p':
        continue
    texts = child.findall('.//w:t', nsmap)
    text = ''.join(t.text or '' for t in texts)

    if '最终得到五张表' in text:
        target_indices['update_176'] = ci
    elif text.strip() == '工资表（gongzi）':
        target_indices['delete_start'] = ci
    elif '规范化为五张表' in text:
        target_indices['update_225'] = ci

# 找到 delete_start 之后的下一个空行 + 第三部分之间的所有元素
# delete_start = [204], delete_end = [224] (最后一个空行)
if target_indices['delete_start']:
    ds = target_indices['delete_start']
    # 从 ds 开始，一直删除到 "第三部分" 之前的元素（即 [225] 规范化总结 之前）
    # 实际要删: [204] 到 [224]，即工资表段落到CREATE VIEW后的空行
    end_idx = None
    for ci in range(ds, len(children)):
        tag = children[ci].tag.split('}')[-1]
        if tag == 'p':
            texts = children[ci].findall('.//w:t', nsmap)
            text = ''.join(t.text or '' for t in texts)
            if '规范化为五张表' in text:
                end_idx = ci - 1  # 删到规范化总结之前
                break
    target_indices['delete_end'] = end_idx

print(f'update_176 idx: {target_indices["update_176"]}')
print(f'delete_start idx: {target_indices["delete_start"]}')
print(f'delete_end idx: {target_indices["delete_end"]}')
print(f'update_225 idx: {target_indices["update_225"]}')

# ==== 第二步：删除工资表相关元素（从后往前删，避免索引偏移）====
if target_indices['delete_start'] is not None and target_indices['delete_end'] is not None:
    ds = target_indices['delete_start']
    de = target_indices['delete_end']
    to_remove = list(range(ds, de + 1))
    print(f'\n要删除的元素索引: {to_remove[0]} ~ {to_remove[-1]} (共{len(to_remove)}个)')

    # 验证删除范围
    for idx in to_remove:
        tag = children[idx].tag.split('}')[-1]
        if tag == 'p':
            texts = children[idx].findall('.//w:t', nsmap)
            text = ''.join(t.text or '' for t in texts)[:80]
            print(f'  删除 [{idx}] <w:p> "{text}"')
        elif tag == 'tbl':
            # table
            rows = children[idx].findall('.//w:tr', nsmap)
            cells = rows[0].findall('.//w:tc', nsmap) if rows else []
            cell_text = ''.join(t.text or '' for t in cells[0].findall('.//w:t', nsmap)) if cells else ''
            print(f'  删除 [{idx}] <w:tbl> rows={len(rows)} first_cell="{cell_text[:50]}"')

    # 从后往前删除
    for idx in reversed(to_remove):
        body.remove(children[idx])

    print(f'\n✅ 已删除 {len(to_remove)} 个元素')

# ==== 第三步：更新文本（需要重新获取 children，因为元素已被删除）====
# 但 Paragraph 对象的 XML 元素还在（只是从 body 移除了），直接通过 paragraph 对象修改

# 找到 P166 段落（"最终得到五张表"）
for p in doc.paragraphs:
    text = p.text
    if '最终得到五张表，以下为完整设计' in text:
        for run in p.runs:
            if '五张表' in run.text:
                run.text = run.text.replace('五张表', '四张表')
                print(f'✅ 已更新: "{run.text[:80]}"')
                break
        break

# 找到 P207 段落（"规范化为五张表"）
for p in doc.paragraphs:
    text = p.text
    if '规范化为五张表' in text:
        for run in p.runs:
            if '五张表' in run.text:
                run.text = run.text.replace(
                    '五张表（工程表、职务表、员工表、工时记录表、工资表）',
                    '四张表（工程表、职务表、员工表、工时记录表）'
                )
                print(f'✅ 已更新: "{run.text[:150]}"')
                break
        break

# ==== 第四步：保存 ====
doc.save(SRC)
print(f'\n✅ 文件已保存: {SRC}')

# ==== 第五步：验证 ====
doc2 = Document(SRC)
found_gongzi = False
for p in doc2.paragraphs:
    if '工资表（gongzi）' in p.text:
        found_gongzi = True
        break
if found_gongzi:
    print('⚠️ 警告：仍存在「工资表（gongzi）」相关文本')
else:
    print('✅ 验证通过：「工资表（gongzi）」已完全移除')

# 检查表格
for ti, tbl in enumerate(doc2.tables):
    rows = len(tbl.rows)
    cols = len(tbl.columns)
    row0_texts = []
    for ci in range(min(cols, 4)):
        row0_texts.append(tbl.cell(0, ci).text.strip()[:30])
    header = ' | '.join(row0_texts)
    if '工资' in header or 'gongzi' in header.lower():
        print(f'⚠️ 警告：Table{ti} 包含工资表内容: {header}')

print(f'\n最终: {len(doc2.paragraphs)} 段落, {len(doc2.tables)} 表格')
