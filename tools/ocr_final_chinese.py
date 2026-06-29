r"""
品牌管理 PPT - 最终版 OCR → 中文 Word 文档
- 支持 JPG + HEIC（ffmpeg 转换）
- 只提取中文内容，过滤英文
- 对比旧版 OCR 缓存，合并优质内容
"""
import subprocess
import os
import re
import sys
import json
from pathlib import Path
from PIL import Image, ImageFilter, ImageEnhance
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ===== 配置 =====
NEW_IMG_DIR = Path(r"C:\Users\AUSU\Desktop\图片2")
CONVERTED_DIR = NEW_IMG_DIR / "converted"
OLD_CACHE_PATH = Path(r"C:\Users\AUSU\Documents\trae_projects\cc\tmp\ocr_cache.json")
NEW_CACHE_PATH = Path(r"C:\Users\AUSU\Documents\trae_projects\cc\tmp\ocr_cache_new.json")
OUTPUT_PATH = Path(r"C:\Users\AUSU\Desktop\品牌管理课程笔记_优化版.docx")
TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSDATA = os.path.expandvars(r"%USERPROFILE%\tesseract-tessdata")
FFMPEG = r"C:\Users\AUSU\AppData\Local\Microsoft\WinGet\Packages\Gyan.FFmpeg_Microsoft.Winget.Source_8wekyb3d8bbwe\ffmpeg-8.1.1-full_build\bin\ffmpeg.exe"


# ===== 图片处理 =====
def convert_heic_to_jpg(heic_path: Path) -> Path:
    """用 ffmpeg 转换 HEIC → JPG"""
    CONVERTED_DIR.mkdir(exist_ok=True)
    jpg_path = CONVERTED_DIR / (heic_path.stem + ".jpg")
    if jpg_path.exists():
        return jpg_path
    r = subprocess.run(
        [FFMPEG, "-y", "-i", str(heic_path), str(jpg_path)],
        capture_output=True, timeout=30
    )
    if r.returncode == 0 and jpg_path.exists():
        return jpg_path
    print(f"  [WARN] HEIC conversion failed: {heic_path.name}")
    return None


def preprocess_image(img: Image.Image) -> Image.Image:
    """图像预处理：缩放 + 灰度 + 增强对比度"""
    w, h = img.size
    # 缩放到 2000px 宽
    if w > 2000:
        scale = 2000 / w
        img = img.resize((2000, int(h * scale)), Image.LANCZOS)
    gray = img.convert("L")
    enhancer = ImageEnhance.Contrast(gray)
    enhanced = enhancer.enhance(2.5)
    enhanced = enhanced.filter(ImageFilter.SHARPEN)
    return enhanced


# ===== OCR =====
def ocr_image(img_path: Path, cache: dict) -> str:
    """OCR 单张图片"""
    cache_key = img_path.name
    if cache_key in cache:
        return cache[cache_key]

    try:
        img = Image.open(img_path)
        processed = preprocess_image(img)
        tmp = NEW_IMG_DIR / f"_ocr_tmp_{img_path.stem}.png"
        processed.save(tmp)

        env = os.environ.copy()
        env["TESSDATA_PREFIX"] = TESSDATA
        r = subprocess.run(
            [TESSERACT, str(tmp), "stdout", "-l", "chi_sim", "--psm", "6"],
            capture_output=True, text=True, env=env, timeout=60, encoding="utf-8"
        )
        tmp.unlink()
        text = r.stdout.strip()
        cache[cache_key] = text
        return text
    except Exception as e:
        print(f"  [ERR] {img_path.name}: {e}")
        return ""


# ===== 中文提取 =====
def chinese_only(text: str) -> list:
    """
    从 OCR 文本中只提取中文内容行。
    规则：
    - 保留含有中文的行
    - 去掉纯英文行
    - 去掉噪声（中文占比太低）
    - 清理 OCR 符号残留
    """
    lines = text.split("\n")
    result = []
    prev_empty = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if not prev_empty:
                result.append("")
                prev_empty = True
            continue
        prev_empty = False

        # 统计中文
        cn_count = sum(1 for c in stripped if '一' <= c <= '鿿')
        en_count = sum(1 for c in stripped if 'a' <= c.lower() <= 'z')
        digit_count = sum(1 for c in stripped if '0' <= c <= '9')
        total = len(stripped)

        # 跳过纯英文/数字行
        if cn_count == 0:
            continue

        # 中文占比太低 = 噪声
        cn_ratio = cn_count / total
        if cn_ratio < 0.15 and cn_count < 3:
            continue

        # 清理行：保留中文、数字、中文标点、常用符号
        cleaned = []
        for c in stripped:
            if ('一' <= c <= '鿿' or          # 中文
                '0' <= c <= '9' or            # 数字
                c in '，。、；：！？…—（）《》〈〉""''【】％％／＋－×＝·①②③④⑤⑥⑦⑧⑨⑩' or
                c in '一二三四五六七八九十百千万亿' or
                c in '．：·' or
                c == ' ' or c == '\t'):
                cleaned.append(c)
        cleaned_line = ''.join(cleaned).strip()

        if cleaned_line and len(cleaned_line) >= 2:
            result.append(cleaned_line)

    # 去首尾空行
    while result and not result[0]:
        result.pop(0)
    while result and not result[-1]:
        result.pop()
    return result


# ===== 标题检测 =====
def looks_like_heading(line: str) -> bool:
    """检测中文标题行"""
    cn = sum(1 for c in line if '一' <= c <= '鿿')
    total = len(line)

    if cn < 3:
        return False
    if cn / total < 0.5:
        return False

    # 序号 + 中文
    if re.match(r'^[一二三四五六七八九十\d]+[、.)．）]\s*[一-鿿]', line):
        return True
    # "第X章/讲"
    if re.match(r'^第[一二三四五六七八九十\d]+[章节讲课]', line):
        return True
    # 短中文行 (3-15字)
    if 3 <= cn <= 15 and cn / total >= 0.6:
        return True
    # 带冒号的标题格式
    if '：' in line:
        parts = line.split('：')
        if len(parts[0]) <= 12 and sum(1 for c in parts[0] if '一' <= c <= '鿿') >= 1:
            return True
    return False


# ===== 对比旧版 =====
def load_old_cache() -> dict:
    """加载旧版 OCR 缓存"""
    if OLD_CACHE_PATH.exists():
        try:
            with open(OLD_CACHE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def compare_quality(new_lines: list, old_text: str) -> str:
    """
    比较新旧 OCR 质量，返回质量标记。
    """
    new_cn = sum(sum(1 for c in l if '一' <= c <= '鿿') for l in new_lines)
    old_cn = sum(1 for c in old_text if '一' <= c <= '鿿')
    if new_cn > old_cn * 1.5:
        return " [+] new better"
    elif old_cn > new_cn * 1.5:
        return " [!] old better"
    else:
        return " [=] similar"


# ===== Word 生成 =====
def create_docx(all_slides: list, doc_path: Path):
    """生成中文 Word 文档"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # === 封面 ===
    title = doc.add_heading("品牌管理课程笔记", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("基于高清 PPT 拍照图片 · OCR 自动识别\n仅提取中文内容 · 优化版").font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"整理日期：2026年6月26日 | {len(all_slides)} 页")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)

    doc.add_page_break()

    # === 正文 ===
    for slide_idx, (fname, cn_lines, quality_note) in enumerate(all_slides):
        if not cn_lines:
            continue

        if slide_idx > 0:
            doc.add_paragraph()

        # 页码
        marker = doc.add_paragraph()
        marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = marker.add_run(f"-- Slide {slide_idx + 1} {quality_note} --")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.italic = True

        # 内容
        for line in cn_lines:
            if not line:
                doc.add_paragraph()
                continue
            if looks_like_heading(line):
                doc.add_heading(line, level=2)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.size = Pt(11)

    # === 尾注 ===
    doc.add_page_break()
    doc.add_heading("关于本笔记", level=1)
    note = doc.add_paragraph()
    note.add_run("整理说明：\n").bold = True
    note.add_run(
        "1. 本笔记基于手机拍摄的高清 PPT 图片，通过 Tesseract OCR 自动识别生成。\n"
        "2. 仅提取中文内容，英文部分已自动过滤。\n"
        "3. OCR 识别可能存在误差，建议对照原始图片核对重要概念和数字。\n"
        "4. 原始图片位于「桌面/图片2」文件夹，按文件名排序对应幻灯片顺序。\n"
        "5. 整理日期：2026年6月26日\n"
    )

    doc.save(str(doc_path))
    print(f"\n  [OK] Document saved: {doc_path}")


# ===== 主流程 =====
def main():
    print("=" * 60)
    print("  Brand Management PPT -> Chinese Word (Final)")
    print("=" * 60)

    # 1. 收集图片（JPG + 转换 HEIC）
    print("\n[Step 1] Collecting images...")
    images = []
    for f in sorted(NEW_IMG_DIR.glob("*")):
        if f.suffix.lower() == '.heic' and f.name.startswith('IMG_'):
            jpg = convert_heic_to_jpg(f)
            if jpg:
                images.append(('HEIC', jpg, f.name))
            else:
                print(f"  [SKIP] Failed to convert: {f.name}")
        elif f.suffix.lower() in ('.jpg', '.jpeg') and f.name.startswith('IMG_'):
            images.append(('JPG', f, f.name))

    print(f"  Total: {len(images)} images ready ({sum(1 for t,_,_ in images if t=='HEIC')} HEIC converted)")

    # 2. 加载缓存
    new_cache = {}
    if NEW_CACHE_PATH.exists():
        try:
            with open(NEW_CACHE_PATH, "r", encoding="utf-8") as f:
                new_cache = json.load(f)
            print(f"  [Cache] {len(new_cache)} existing entries loaded")
        except Exception:
            pass

    old_cache = load_old_cache()
    print(f"  [Old] {len(old_cache)} old entries for comparison")

    # 3. 批量 OCR
    print(f"\n[Step 2] OCR processing {len(images)} images...")
    all_slides = []

    for idx, (img_type, img_path, orig_name) in enumerate(images, 1):
        cached = " (cached)" if orig_name in new_cache else ""
        print(f"  [{idx:2d}/{len(images)}] {orig_name}{cached} ...", end=" ", flush=True)
        text = ocr_image(img_path, new_cache)
        cn_lines = chinese_only(text)
        cn_chars = sum(len(l) for l in cn_lines)
        print(f"{cn_chars:4d} cn chars, {len(cn_lines):2d} lines", end="")

        # 对比旧版
        quality = ""
        if old_cache:
            # 简单取旧版前几个结果对比
            old_vals = list(old_cache.values())
            if idx <= len(old_vals):
                quality = compare_quality(cn_lines, old_vals[idx - 1])
        print(quality)

        all_slides.append((orig_name, cn_lines, quality))

        if idx % 15 == 0:
            with open(NEW_CACHE_PATH, "w", encoding="utf-8") as f:
                json.dump(new_cache, f, ensure_ascii=False)

    # 保存缓存
    with open(NEW_CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(new_cache, f, ensure_ascii=False)

    # 统计
    total_cn = sum(sum(len(l) for l in slide[1]) for slide in all_slides)
    valid_slides = sum(1 for s in all_slides if s[1])
    print(f"\n[Stats] {valid_slides}/{len(images)} slides with CN text, ~{total_cn:,} CN chars")

    # 4. 生成 Word
    print(f"\n[Step 3] Generating Word document...")
    create_docx(all_slides, OUTPUT_PATH)

    # 清理
    print(f"\n[Clean] Removing temp files...")
    for tmp in list(NEW_IMG_DIR.glob("_ocr_tmp_*.png")):
        try: tmp.unlink()
        except: pass

    print(f"\n{'=' * 60}")
    print(f"  [DONE] Complete!")
    print(f"  Output: {OUTPUT_PATH}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
