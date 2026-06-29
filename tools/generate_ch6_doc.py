# -*- coding: utf-8 -*-
"""
生成《第六章 智慧库存管理》知识点总结 + 案例分析 Word 文档
"""
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

doc = Document()

# ── 全局样式设置 ──
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
# 设置中文字体
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ── 辅助函数 ──
def add_heading_styled(text, level=1):
    """添加带样式的标题"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '微软雅黑'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return heading

def add_para(text, bold=False, indent=False):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(11)
    run.bold = bold
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_bullet(text, level=0):
    """添加要点（带圆点符号）"""
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = '微软雅黑'
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * (level + 1))
    return p

def add_key_point(label, content):
    """添加【知识点】段落"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    run_label = p.add_run(f'【{label}】')
    run_label.bold = True
    run_label.font.name = '微软雅黑'
    run_label.font.size = Pt(11)
    run_label.font.color.rgb = RGBColor(0, 112, 192)
    run_label.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_text = p.add_run(content)
    run_text.font.name = '微软雅黑'
    run_text.font.size = Pt(11)
    run_text.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# ============================================================
# 封面
# ============================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第六章  智慧库存管理')
run.font.name = '微软雅黑'
run.font.size = Pt(26)
run.bold = True
run.font.color.rgb = RGBColor(0, 112, 192)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('章节知识点总结 · 开篇案例 · 章末案例分析')
run.font.name = '微软雅黑'
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(89, 89, 89)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run('智慧供应链管理课程作业')
run.font.name = '微软雅黑'
run.font.size = Pt(12)
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_page_break()

# ============================================================
# 目录概要
# ============================================================
add_heading_styled('本章结构', level=1)
add_para('第六章围绕"智慧库存管理"展开，共分为五个部分，从概念到应用层层递进：')

toc_items = [
    '6.1  智慧库存概述 —— 定义、目标与意义',
    '6.2  传统库存管理存在的问题 —— 三大核心痛点',
    '6.3  智慧库存系统的机理与架构 —— 预测、仓储、运输三大系统',
    '6.4  智慧库存管理信息平台 —— 设计原则、功能体系、架构分层',
    '6.5  智慧库存管理的组织与应用 —— 多层面实施 + 日日顺案例',
]
for item in toc_items:
    add_bullet(item)

doc.add_page_break()

# ============================================================
# 6.1 智慧库存概述
# ============================================================
add_heading_styled('6.1  智慧库存概述', level=1)

add_heading_styled('一、库存的定义与分类', level=2)
add_para('库存包括原材料、半成品、成品等，是企业运营的关键组成部分。', indent=True)

add_heading_styled('二、智慧库存管理的目标', level=2)
goals = [
    ('高效响应与服务', '实时监控库存，快速响应需求，提供高效、个性化服务。'),
    ('防止过量存货与生产瓶颈', '精准预测和监控防止库存积压，避免生产瓶颈和资金占用。'),
    ('降低成本与提高竞争力', '降低库存成本，提高周转率，增强企业竞争力。'),
]
for title, desc in goals:
    add_key_point(title, desc)

add_heading_styled('三、智慧库存管理的意义', level=2)
meanings = [
    '提高库存效率：通过智能算法和数据分析，准确预测需求，优化库存水平。',
    '减少供应链风险：实时监控库存和供应链信息，及时发现和应对风险。',
    '提升客户满意度：加快订单处理和交付速度，提供准确的库存信息。',
    '优化供应链协调与合作：促进供应链各环节的协同工作，提高整体效率。',
    '推动企业创新和竞争力：为企业提供数据支持和决策依据，推动创新发展。',
]
for m in meanings:
    add_bullet(m)

doc.add_page_break()

# ============================================================
# 6.2 传统库存管理存在的问题
# ============================================================
add_heading_styled('6.2  传统库存管理存在的问题', level=1)

add_para('传统库存管理面临三大核心痛点，这些问题制约了供应链效率的提升：', indent=True)

add_heading_styled('痛点一：需求预测不够准确', level=2)
add_bullet('基于历史数据的局限：市场需求动态变化，历史数据难以准确预测未来需求。')
add_bullet('缺乏实时数据：无法及时了解市场变化和客户需求的变化。')
add_bullet('人工主观干预：依赖人的主观判断和经验，容易出现误差。')
add_bullet('缺乏全局视野：以单个产品或环节为基础，难以考虑整个供应链的协同。')

add_heading_styled('痛点二：库存信息共享不畅', level=2)
add_bullet('信息流通速度慢：供应链各环节之间信息传递不及时，导致决策滞后。')
add_bullet('信息不对称：不同环节之间的信息不共享，难以形成整体协同。')

add_heading_styled('痛点三：缺乏灵活性和响应能力', level=2)
add_bullet('批量生产和订购：难以适应市场需求的快速变化。')
add_bullet('刚性供应链：调整产品组合和生产计划困难，导致库存积压和过时产品。')

doc.add_page_break()

# ============================================================
# 6.3 智慧库存系统的机理与架构
# ============================================================
add_heading_styled('6.3  智慧库存系统的机理与架构', level=1)

# --- 6.3.1 智慧预测系统 ---
add_heading_styled('6.3.1  智慧预测系统', level=2)

add_para('（一）核心要素与特征', bold=True)
pred_features = [
    '数据采集和整合：整合多源数据，进行多维分析。',
    '需求预测和优化：考虑多种因素，进行精确预测和优化。',
    '库存管理和优化：优化多级库存，确定补货策略。',
    '供应链协同和合作：与供应链伙伴共享信息，实现协同决策。',
    '产能规划和优化：识别和管理产能约束，提高产能利用率。',
    '数据分析和决策支持：提供高级分析功能，支持决策制定。',
]
for f in pred_features:
    add_bullet(f)

add_para('（二）典型应用实践', bold=True)
add_bullet('欧睿数据的应用：为时尚企业提供商品管理解决方案，实现精准决策。')
add_bullet('阿里巴巴供应链中台的应用：实现智能预测备货、选品和分仓调拨，提高供应链效率。')

# --- 6.3.2 智慧仓储系统 ---
add_heading_styled('6.3.2  智慧仓储系统', level=2)

add_para('（一）六大功能模块', bold=True)
warehouse_modules = [
    '仓库物流管理模块：实现仓库作业的信息化管理。',
    '仓库设备管理模块：科学调度和管理库内智能设备。',
    '订单管理模块：自动完成订单的采集和处理。',
    '财务信息管理模块：生成财务报表，实现财务管理自动化。',
    '客户管理模块：智能分析客户需求，提供个性化服务。',
    '人力资源管理模块：动态调度作业人员，提高人力资源利用率。',
]
for m in warehouse_modules:
    add_bullet(m)

add_para('（二）技术架构', bold=True)
add_bullet('分为技术层、控制层和业务层，实现智能化管理。')

add_para('（三）典型应用实践', bold=True)
add_bullet('自动导引车（AGV）：实现物流搬运的自动化，提高工作效率。')
add_bullet('码垛机器人：完成物料的码垛和搬运任务，提高作业精度。')
add_bullet('自动分拣机：实现货物的快速分拣，提高配送效率。')

# --- 6.3.3 智慧运输系统 ---
add_heading_styled('6.3.3  智慧运输系统', level=2)

add_para('（一）核心要素与特征', bold=True)
transport_features = [
    '实时数据采集与共享：通过传感器和物联网技术，实时获取运输数据。',
    '智能调度与路径优化：根据实时数据，优化运输路线和调度方案。',
    '车辆追踪与监控：实时跟踪车辆位置和状态，确保运输安全。',
    '货物追踪与管理：实现货物的全程追踪，保证货物安全。',
    '数据分析与预测：利用大数据分析，预测运输需求和市场趋势。',
    '资源协同与共享：促进运输资源的协同和共享，提高资源利用率。',
    '预警与风险管理：及时发现和预警潜在风险，降低损失。',
]
for f in transport_features:
    add_bullet(f)

add_para('（二）功能应用场景', bold=True)
add_bullet('电子不停车收费（ETC）：提高交通基础设施的运行效率。')
add_bullet('智能终端应用：为公众提供便捷的出行服务。')
add_bullet('运载工具跟踪监管：改善运输安全，提高运输效率。')
add_bullet('货运组织管理：优化货运组织，提高物流效益。')
add_bullet('物流跟踪与可视化：实现物流过程的透明化管理。')

doc.add_page_break()

# ============================================================
# 6.4 智慧库存管理信息平台
# ============================================================
add_heading_styled('6.4  智慧库存管理信息平台', level=1)

add_heading_styled('6.4.1  设计原则', level=2)
principles = [
    ('实时性与准确性', '及时获取和更新库存数据，确保数据的准确性。'),
    ('弹性与可扩展性', '适应业务变化，满足企业不断发展的需求。'),
    ('预测与优化', '具备预测和优化库存的能力，提高库存管理水平。'),
    ('可视化与易用性', '提供直观的界面，方便用户操作和管理。'),
    ('整合与协同', '整合供应链各环节的数据，实现协同管理。'),
    ('安全与隐私', '保障库存数据的安全和隐私。'),
]
for title, desc in principles:
    add_key_point(f'原则{principles.index((title,desc))+1}', f'{title}：{desc}')

add_heading_styled('6.4.2  设计目标', level=2)
goals_platform = [
    '自动化采购执行：实现采购流程的自动化，提高工作效率。',
    '前瞻性市场洞察：提供市场洞察和分析，支持决策制定。',
    '可视化采购控制：提供采购控制塔，实现采购过程的可视和可控。',
    '数字化协同网络：构建供需双方的协同网络，实现数字化协同。',
]
for g in goals_platform:
    add_bullet(g)

add_heading_styled('6.4.3  功能体系与设计架构', level=2)

add_para('（一）核心功能模块', bold=True)
add_bullet('库存监控：实时监控库存数量、位置和状态，提供库存报表和分析功能。')
add_bullet('库存管理与操作：支持入库、出库、调拨等操作，实现精细化管理。')
add_bullet('库存需求预测与计划：根据历史数据和市场趋势，预测库存需求，制定合理计划。')
add_bullet('库存风险管理：监测和识别库存风险，提供预警机制和应对措施。')
add_bullet('用户权限与安全管理：设定用户权限，保障库存数据安全和隐私。')

add_para('（二）四层架构设计', bold=True)
add_bullet('数据采集与存储层：负责采集和存储库存相关数据，确保数据质量和完整性。')
add_bullet('数据处理与分析层：对采集到的数据进行处理和分析，提取有价值的信息。')
add_bullet('应用与业务逻辑层：实现库存管理的各项功能（库存查询、入库出库管理等）。')
add_bullet('用户界面层：提供友好的用户界面，方便用户操作和管理库存。')

doc.add_page_break()

# ============================================================
# 6.5 智慧库存管理的组织与应用
# ============================================================
add_heading_styled('6.5  智慧库存管理的组织与应用', level=1)

add_heading_styled('6.5.1  多层面组织实施', level=2)

add_para('智慧库存管理的实施需要国家、行业、企业三个层面协同推进：', indent=True)

add_para('（一）国家层面', bold=True)
add_bullet('政策法规支持：出台相关政策法规，鼓励和支持智慧库存管理的发展。')
add_bullet('平台建设与数据标准：建设国家级智慧库存管理平台，制定统一的数据标准。')
add_bullet('技术创新与人才培养：鼓励技术创新，培养智慧库存管理的专业人才。')

add_para('（二）行业层面', bold=True)
add_bullet('制定标准与规范：制定行业智慧库存管理的标准和规范。')
add_bullet('建立评估和奖励机制：激励企业积极推进智慧库存管理。')
add_bullet('促进数据共享与分析：实现行业内数据的共享和分析，提高行业整体水平。')

add_para('（三）企业层面', bold=True)
add_bullet('评估现状与确定目标：评估当前库存管理状况，确定智慧库存管理目标。')
add_bullet('选择解决方案：选择适合企业的智慧库存管理系统和技术。')
add_bullet('系统实施与数据管理：实施智慧库存管理系统，整理和管理库存数据。')
add_bullet('人员培训与协作：培训员工，促进各部门之间的协作和沟通。')

add_heading_styled('6.5.2  智慧仓储应用形态', level=2)

add_para('（一）无人仓', bold=True)
add_bullet('概念：实现仓库作业的无人化，包括货物搬运、存储、拣选等。')
add_bullet('主要构成：由搬运设备、存储设备、拣选设备等组成，通过人工智能算法实现自主决策和控制。')

add_para('（二）智慧云仓（以京东为例）', bold=True)
add_bullet('技术装备：包括地狼、天狼、无人叉车等机器人，实现仓储作业的自动化。')
add_bullet('应用效果：提高仓储效率，降低成本，为物流行业的发展提供新的动力。')

doc.add_page_break()

# ============================================================
# 开篇案例：京东智慧供应链
# ============================================================
add_heading_styled('开篇案例：京东推出"智慧供应链"', level=1)

add_para('【案例摘要】', bold=True)
add_para('2016年5月，京东成立了智慧物流开放平台"X事业部"，自主研发无人机、无人车、无人仓、无人配送站等一系列尖端智慧物流项目。2017年10月，京东物流首个全流程无人仓正式亮相上海——这是全球首个正式落成并规模化投入使用的全流程无人物流中心。', indent=True)
add_para('京东无人仓建筑面积40,000m²，由收货、存储、订单拣选、包装4个作业系统组成。其智慧大脑能在0.2秒内计算出300多个机器人运行的680亿条可行路径；智能控制系统反应速度是人的6倍；分拣速度达3m/s，为全世界最快；运营效率是传统仓库的10倍。', indent=True)

add_para('【核心知识点提炼】', bold=True)

add_key_point('智能预测', '京东的预测技术是智慧供应链的"大脑"——通过大数据、采销经验和算法，预先知道消费者行为（买什么、买多少、以什么价格买），实现需求预测。')
add_key_point('智能补货与内配', '在京东全国的300多个仓储之间做到内配，提前将商品运到离消费者最近的仓库，缩短配送时间。')
add_key_point('库存周转率', '京东自营商品超900万SKU，库存周转仅需31天。对比Costco管理数千SKU约30天，京东在SKU规模远超的情况下实现了同等的周转效率。')
add_key_point('四个维度', '数据挖掘、人工智能、流程再造和技术驱动。')
add_key_point('五大层面', '覆盖商品、价格、计划、库存、协同五大层面的智慧供应链解决方案。')

doc.add_page_break()

# ============================================================
# 章末案例：日日顺物流
# ============================================================
add_heading_styled('章末案例：日日顺物流智能仓', level=1)

add_para('【案例摘要】', bold=True)
add_para('日日顺物流是海尔集团旗下的大件物流领导企业，率先建成国内首个大件物流智能仓——黄岛智能仓。1.8万m²的智能仓内，AGV、RGV、堆垛机、输送机等智能设备无缝协同运转，由大数据总控系统作为"智慧大脑"统一调度，实现了"零搬运、零差错、零货损"。效率提升约1倍，人工成本降低50%。', indent=True)
add_para('日日顺物流已由传统物流企业转变为物联网时代开放的场景物流生态平台，拥有136个智慧仓、6000多家微仓、3300多条班车循环专线，通过"车小微"平台连接20万服务人员，为用户提供全流程个性化服务方案。', indent=True)

add_heading_styled('案例分析：三个问题', level=2)

# --- 问题1 ---
add_heading_styled('问题1：日日顺物流的智能仓储的优势体现在哪些方面？', level=3)

add_para('日日顺智能仓储的优势体现在技术、效率、服务和协同四个层面：')

add_key_point('自动化与智能化设备应用', 'AGV、RGV（有轨制导车）、堆垛机、输送机等针对大件商品定制的多种机器人无缝协同运转，体现了智能装备集成能力。')
add_key_point('数据驱动的"智慧大脑"', '大数据总控系统集数据中心、监控中心、决策中心和控制中心于一体，实现算法指导生产。')
add_key_point('卓越的运营绩效', '作业效率提升约1倍，人工成本降低50%，实现"零搬运、零差错、零货损"——体现了精益物流与运营效率优化。')
add_key_point('端到端的供应链协同', '智能仓前端对接产业（原材料即需即供），后端服务用户（商品流转<24小时），实现了一体化供应链协同。')
add_key_point('个性化增值服务', '在提供送装服务同时，还可提供除甲醛、净化空气、家电保养等生活场景服务方案。')

# --- 问题2 ---
add_heading_styled('问题2：限制大件物流智能仓发展的因素是什么？', level=3)

add_para('限制大件物流智能仓发展的核心因素源于大件商品自身的物理特性：')

add_key_point('商品物理属性的制约', '大件商品体积大、质量重、形状不规则，使得通用的小件物流自动化设备（如标准分拣机、小型AGV）无法直接适用，必须进行高成本的定制化开发。这是行业发展的主要瓶颈。')
add_key_point('技术适配难度高', '大件物流的搬运、存储、拣选等环节对设备的承重能力、精度控制和稳定性要求远高于小件物流，技术门槛更高。')
add_key_point('投资回报周期长', '大件智能仓的建设和设备定制成本高昂，而大件商品周转速度相对较慢，投资回报周期较长，抑制了企业投入意愿。')

# --- 问题3 ---
add_heading_styled('问题3：日日顺物流如何打造场景化物流生态？', level=3)

add_para('日日顺通过从"物流执行者"向"场景服务商"的战略跃迁，构建了独特的物流生态：')

add_key_point('服务内容的场景化延伸', '将物流触点转化为生活场景服务入口——在送装家电时，可同步提供除甲醛、净化空气、后续保养等服务方案，实现从"送货"到"送生活方案"的升级。')
add_key_point('资源整合与平台化运营', '搭建开放的"车小微"平台，整合社会化闲散汽车和司机资源，形成覆盖全国的触点网络。服务范围从送装扩展到同城用车、售后维修、代发快递、产品推荐等。')
add_key_point('商业模式的根本转型', '从物流企业转变为物联网时代开放的场景物流生态平台：从商品配送 → 终身用户交互；从交易平台 → 生态平台。核心竞争力从内部资产转向终身用户数量。')
add_key_point('五大场景生态', '打造了美好住居场景、健康生活场景、三农服务场景、环保出行场景、产业互联场景等，为用户提供全流程个性化方案。')

doc.add_page_break()

# ============================================================
# 战略与管理分析
# ============================================================
add_heading_styled('日日顺案例深度分析', level=1)

add_heading_styled('一、战略层面', level=2)
add_para('日日顺物流的战略核心是生态品牌战略。它不再局限于赚取物流运费的交易型思维，而是以用户为中心，围绕用户的实际生活场景（如"美好住居"、"健康生活"）来创造和传递价值。通过"车小微"平台，将20万服务兵变为与用户零距离交互的触点，目标是获取并经营"终身用户"。')

add_heading_styled('二、管理层面', level=2)
add_para('日日顺践行了海尔首创的"人单合一"模式。该模式打破了传统科层制，让每一位"车小微"创客都直接面对用户需求，拥有决策权、用人权和分配权，从而能快速响应并满足用户的个性化、场景化需求。')

add_heading_styled('三、运营层面', level=2)
add_para('日日顺构建了三级智慧物流网络（136个智慧仓、6000多家微仓、3300多条专线），实现全国范围内的高效覆盖。黄岛智能仓作为核心节点，通过大数据总控系统实现对全网资源的实时监控与优化调度，确保从前端工厂柔性生产（即需即供）到后端用户极速交付（<24小时）的高效协同。')

add_heading_styled('四、技术层面', level=2)
add_para('硬科技方面，成功将AGV、RGV、堆垛机等自动化设备首次规模化应用于大件物流领域。软实力方面，构建了以大数据总控系统为核心的智慧大脑，实现了对海量运营数据的采集、分析、决策与控制——真正做到了"算法指导生产"。')

add_heading_styled('五、结论', level=2)
add_para('在物联网时代，物流企业的竞争已超越单纯的运输与仓储效率，转向基于场景的生态竞争。日日顺通过战略上的生态化转型、管理上的"人单合一"模式、运营上的全国网络协同以及技术上的软硬结合，成功破解了大件物流智能化的难题，重新定义了物流的价值——从商品的位移工具升级为连接用户与产业、创造美好生活的场景服务平台。')

doc.add_page_break()

# ============================================================
# 课后习题
# ============================================================
add_heading_styled('课后习题（参考答案要点）', level=1)

add_heading_styled('1. 传统库存存在哪些问题？', level=2)
add_para('传统库存管理存在三大核心问题：')
add_bullet('需求预测不够准确——依赖历史数据、缺乏实时数据、人工主观干预、缺乏全局视野。')
add_bullet('库存信息共享不畅——信息流通速度慢、信息不对称、难以形成整体协同。')
add_bullet('缺乏灵活性和响应能力——批量生产和订购模式难以适应市场快速变化，刚性供应链导致库存积压。')

add_heading_styled('2. 智慧库存系统的机理与架构是什么？', level=2)
add_para('智慧库存系统由三大子系统构成：')
add_bullet('智慧预测系统：整合多源数据→需求预测与优化→库存管理与优化→供应链协同→产能规划→数据分析与决策支持。')
add_bullet('智慧仓储系统：六大功能模块（物流管理/设备管理/订单管理/财务信息/客户管理/人力资源）+ 三层技术架构（技术层/控制层/业务层）+ 智能设备（AGV/码垛机器人/自动分拣机）。')
add_bullet('智慧运输系统：实时数据采集→智能调度与路径优化→车辆/货物追踪→数据分析与预测→资源协同→预警管理。')

add_heading_styled('3. 智慧库存管理信息平台设计架构是什么？', level=2)
add_para('采用四层架构设计：')
add_bullet('数据采集与存储层 → 数据处理与分析层 → 应用与业务逻辑层 → 用户界面层')
add_para('同时遵循六大设计原则（实时准确/弹性可扩展/预测优化/可视化易用/整合协同/安全隐私）和四大设计目标（自动化采购/前瞻洞察/可视化控制/数字化协同）。')

# ============================================================
# 保存文档
# ============================================================
output_path = r'C:\Users\AUSU\Desktop\第六章_智慧库存管理_知识点与案例总结.docx'
doc.save(output_path)
print(f'文档已保存至：{output_path}')
