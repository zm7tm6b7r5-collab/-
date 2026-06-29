"""
Anthropic 品牌美学 — 全文档配色与字体
参考: github.com/anthropics/skills — brand-guidelines
中文适配: Poppins→黑体 Lora→宋体 (仅修改格式，不动内容)
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re, os

SRC = r"C:\Users\AUSU\Desktop\数据库原理期末大作业(3).docx"
DST = r"C:\Users\AUSU\Desktop\数据库原理期末大作业(3)_美化版.docx"

# ══════════════════════════════════════════
# Anthropic 品牌色
# ══════════════════════════════════════════
INK    = RGBColor(0x14, 0x14, 0x13)   # 暖黑 — 正文/二级标题
CORAL  = RGBColor(0xD9, 0x77, 0x57)   # 陶土橙 — 主标题/一级标题 (标志色)
GREEN  = RGBColor(0x78, 0x8C, 0x5D)   # 橄榄绿 — 表注/图注
CANVAS = RGBColor(0xFA, 0xF9, 0xF5)   # 暖奶油白 — 表格文字/交替行底色
CANVAS_HEX = "FAF9F5"                  # 同上（shade用）
CORAL_HEX  = "D97757"                  # 表头底色

# ══════════════════════════════════════════
# 字体 (Anthropic 中文适配)
#   Poppins → 黑体   Lora → 宋体
# ══════════════════════════════════════════
FONT_HEADING = '黑体'
FONT_BODY    = '宋体'
FONT_CODE    = 'Consolas'

SIZE_TITLE = Pt(22)
SIZE_H1    = Pt(16)
SIZE_H2    = Pt(14)
SIZE_BODY  = Pt(12)
SIZE_CODE  = Pt(10)
SIZE_TABLE = Pt(10.5)

# SQL 代码识别
SQL_STARTS = ('SELECT ', 'INSERT ', 'UPDATE ', 'DELETE ', 'CREATE ',
              'DROP ', 'ALTER ', 'USE ', 'SET ', '--')
SQL_KW = ['SELECT','FROM','WHERE','INSERT','CREATE','TABLE','PRIMARY',
          'KEY','FOREIGN','REFERENCES','ENGINE','INTO','VALUES','JOIN',
          'LEFT','RIGHT','INNER','ON','AND','OR','GROUP','BY','ORDER',
          'COUNT','MAX','MIN','AS','DISTINCT','LIMIT','LIKE','IN',
          'BETWEEN','CASE','WHEN','THEN','ELSE','END','UNION','HAVING',
          'DROP','ALTER','CHAR','VARCHAR','INT','TEXT','DATE','FLOAT',
          'DECIMAL','CHARSET','utf8mb4','COLLATE','CASCADE','INCREMENT',
          'DATABASE','IF','CURDATE','YEAR','GROUP_CONCAT','SEPARATOR',
          'EXISTS','NOT','NULL','DEFAULT','SET']

def is_code(para):
    text = para.text.strip()
    if not text: return False
    upper = text.upper()
    for kw in SQL_STARTS:
        if upper.startswith(kw): return True
    if re.search(r'[一-鿿]', text): return False
    return sum(1 for kw in SQL_KW if kw in upper) >= 2

def is_caption(para):
    t = para.text.strip()
    return bool(re.match(r'[表图]\d+[-–]\d+', t))

def shade_cell(cell, color):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_font(run, name, size, color, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = size
    run.font.color.rgb = color
    run.bold = bold

# ══════════════════════════════════════════
print("[1/3] Reading...")
doc = Document(SRC)

for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.54)
    s.left_margin = s.right_margin = Cm(3.17)

# Normal 默认样式
ns = doc.styles['Normal']
ns.font.name = FONT_BODY
ns.font.size = SIZE_BODY
ns.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_BODY)
ns.paragraph_format.line_spacing = 1.5
ns.paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════
print("[2/3] Styling...")
for para in doc.paragraphs:
    sn = para.style.name if para.style else ''
    text = para.text.strip()
    if not text: continue

    # ── 跳过 SQL 代码 ──
    if is_code(para): continue

    if sn == 'Title':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(36)
        para.paragraph_format.space_after  = Pt(24)
        para.paragraph_format.line_spacing = 1.2
        for r in para.runs:
            set_font(r, FONT_HEADING, SIZE_TITLE, CORAL, bold=False)

    elif sn.startswith('Heading 1'):
        para.paragraph_format.space_before = Pt(28)
        para.paragraph_format.space_after  = Pt(14)
        para.paragraph_format.line_spacing = 1.3
        for r in para.runs:
            set_font(r, FONT_HEADING, SIZE_H1, CORAL, bold=False)

    elif sn.startswith('Heading 2'):
        para.paragraph_format.space_before = Pt(20)
        para.paragraph_format.space_after  = Pt(10)
        para.paragraph_format.line_spacing = 1.3
        for r in para.runs:
            set_font(r, FONT_HEADING, SIZE_H2, INK, bold=False)

    elif is_caption(para):
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(4)
        para.paragraph_format.line_spacing = 1.2
        for r in para.runs:
            set_font(r, FONT_HEADING, Pt(10), GREEN, bold=False)

    else:
        # 正文
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(8)
        for r in para.runs:
            set_font(r, FONT_BODY, SIZE_BODY, INK)

# ══════════════════════════════════════════
# 表格 — Anthropic 配色
# ══════════════════════════════════════════
for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(table.rows):
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.line_spacing = 1.1
                for r in p.runs:
                    if ri == 0:
                        set_font(r, FONT_HEADING, SIZE_TABLE, CANVAS, bold=False)
                    else:
                        set_font(r, FONT_BODY, SIZE_TABLE, INK)
    # 表头 Coral 底 + 交替行 Canvas 底
    if table.rows:
        for c in table.rows[0].cells:
            shade_cell(c, CORAL_HEX)
    for ri in range(1, len(table.rows)):
        if ri % 2 == 0:
            for c in table.rows[ri].cells:
                shade_cell(c, CANVAS_HEX)

# ══════════════════════════════════════════
if os.path.exists(DST):
    os.remove(DST)
print(f"[3/3] Saving...")
doc.save(DST)
print("Done!")
