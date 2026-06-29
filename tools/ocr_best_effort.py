r"""
最后尝试：基于旧版 OCR 缓存（英文质量尚可），尽最大努力提取中文内容。
策略：
1. 使用旧版缓存（对中英文混合识别相对较好）
2. 严格中文提取 + 质量评分
3. 对质量差的页面标记为"需人工核对"
4. 产出草稿级文档
"""
import json
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

OLD_CACHE = Path(r"C:\Users\AUSU\Documents\trae_projects\cc\tmp\ocr_cache.json")
OUTPUT = Path(r"C:\Users\AUSU\Desktop\品牌管理课程笔记_优化版.docx")


def is_real_chinese_word(text: str) -> bool:
    """
    检查文本是否包含真实的常用中文词（而非随机字符）
    """
    # 常见中文二字词模式（社会学/管理学相关）
    common_patterns = [
        '品牌', '管理', '产品', '市场', '营销', '企业', '战略', '定位',
        '消费者', '价值', '竞争', '创新', '发展', '分析', '理论', '概念',
        '特性', '特征', '成功', '组织', '服务', '设计', '沟通', '领导',
        '能力', '知识', '资源', '技能', '危机', '形象', '传播', '策略',
        '方法', '步骤', '案例', '什么是', '如何', '为什么', '分类',
        '关系', '文化', '目标', '需求', '行为', '决策', '体验', '质量',
        '多元化', '差异化', '国际化', '标准化', '个性化',
        '品牌定位', '品牌形象', '品牌价值', '品牌管理', '品牌战略',
        '核心竞争力', '目标市场', '市场份额', '产品生命周期',
        '第一', '第二', '第三', '第四', '首先', '其次', '最后',
        '定义', '特点', '作用', '意义', '影响', '因素',
        '第', '章', '节', '讲', '课',
    ]

    text_clean = text.replace(' ', '').replace('\n', '')
    for pat in common_patterns:
        if pat in text_clean:
            return True
    return False


def chinese_filter_quality(text: str) -> tuple:
    """
    从 OCR 文本提取中文，返回 (行列表, 质量评分 0-100)
    """
    lines = text.split('\n')
    result = []
    total_cn = 0
    good_cn = 0

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 统计中文
        cn_count = sum(1 for c in stripped if '一' <= c <= '鿿')
        total = len(stripped)
        if cn_count == 0:
            continue

        cn_ratio = cn_count / total

        # 质量检查
        is_good = is_real_chinese_word(stripped) and cn_ratio > 0.3

        if is_good:
            # 高质量行：清理后保留
            cleaned = []
            for c in stripped:
                if ('一' <= c <= '鿿' or
                    '0' <= c <= '9' or
                    c in '，。、；：！？…—（）《》【】／＋－＝·①②③④⑤⑥⑦⑧⑨⑩％％' or
                    c in '一二三四五六七八九十百千万亿' or
                    c in '．：· '):
                    cleaned.append(c)
            clean_line = ''.join(cleaned).strip()
            if clean_line and len(clean_line) >= 3:
                result.append(('good', clean_line))
            good_cn += cn_count
        elif cn_ratio > 0.2 and cn_count >= 3:
            # 中等质量：保留但标记
            result.append(('noisy', stripped[:200]))
        total_cn += cn_count

    # 质量评分
    quality = min(100, int((good_cn / max(total_cn, 1)) * 100))
    return result, quality


def looks_like_heading(line: str) -> bool:
    """检测标题"""
    cn = sum(1 for c in line if '一' <= c <= '鿿')
    if cn < 2:
        return False
    if cn / max(len(line), 1) < 0.4:
        return False
    if re.match(r'^[一二三四五六七八九十\d]+[、.)．）]', line):
        return True
    if re.match(r'^第[一二三四五六七八九十\d]+[章节讲课]', line):
        return True
    if 2 <= cn <= 20 and cn / max(len(line), 1) >= 0.5:
        return True
    return False


def create_docx(slides_data: list):
    """生成文档"""
    doc = Document()

    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # 封面
    title = doc.add_heading("品牌管理课程笔记", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("OCR 自动识别 · 草稿版\n⚠️ 仅提取中文内容，准确率有限，建议对照原图核对").font.size = Pt(10)

    doc.add_page_break()

    # 内容
    good_slides = 0
    for fname, lines, quality in slides_data:
        if not lines:
            continue

        quality_label = "高质量" if quality >= 50 else ("中等" if quality >= 20 else "需核对")
        has_good = any(tag == 'good' for tag, _ in lines)

        if has_good:
            good_slides += 1

        # 页标记
        marker = doc.add_paragraph()
        marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = marker.add_run(f"-- {quality_label} (Q{quality}) --")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180) if quality >= 30 else RGBColor(220, 100, 100)
        run.font.italic = True

        for tag, line in lines:
            if looks_like_heading(line):
                h = doc.add_heading(line, level=2)
                if tag == 'noisy':
                    for run in h.runs:
                        run.font.color.rgb = RGBColor(200, 100, 100)
            else:
                p = doc.add_paragraph()
                run = p.add_run(line)
                if tag == 'noisy':
                    run.font.color.rgb = RGBColor(200, 100, 100)
                    run.font.size = Pt(10)
                else:
                    run.font.size = Pt(11)

    doc.add_page_break()
    doc.add_heading("使用说明", level=1)
    doc.add_paragraph(
        "本笔记由 OCR 自动识别生成，准确率有限。\n\n"
        "红色/灰色标记的内容识别质量较低，建议重点核对。\n"
        "黑色文字是系统判定质量较高的内容。\n\n"
        "推荐后续优化方案：\n"
        "1. 使用扫描类 App（如 CamScanner、Microsoft Lens）重新拍摄，再做 OCR\n"
        "2. 使用商业 OCR API（百度 OCR 约 0.01元/张，中文识别准确率 >95%）\n"
        "3. 如能获取 PPT 原文件（.pptx），可直接提取文字\n\n"
        f"统计：{good_slides} 页有较高质量中文内容"
    )

    doc.save(str(OUTPUT))
    print(f"Document saved: {OUTPUT}")
    print(f"Good slides: {good_slides}/{len(slides_data)}")


def main():
    print("Loading old OCR cache...")
    with open(OLD_CACHE, "r", encoding="utf-8") as f:
        cache = json.load(f)

    # Sort by original filename order
    sorted_keys = sorted(cache.keys())
    print(f"Total entries: {len(sorted_keys)}")

    slides_data = []
    for key in sorted_keys:
        text = cache[key]
        lines, quality = chinese_filter_quality(text)
        slides_data.append((key, lines, quality))

    # Stats
    total_good = sum(1 for _, _, q in slides_data if q >= 50)
    total_medium = sum(1 for _, _, q in slides_data if 20 <= q < 50)
    total_bad = sum(1 for _, _, q in slides_data if q < 20)
    print(f"Quality: {total_good} good, {total_medium} medium, {total_bad} poor")

    print("Generating document...")
    create_docx(slides_data)

    print("Done!")


if __name__ == "__main__":
    main()
