# -*- coding: utf-8 -*-
"""
修改数据库原理期末大作业 docx：
替换第一部分（PetStore ER图 + 逻辑结构）为图文表格式呈现
"""
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

SRC = r'C:\Users\AUSU\Desktop\数据库原理期末大作业(3).docx'
DST = r'C:\Users\AUSU\Desktop\数据库原理期末大作业(3).docx'
ER_IMG = r'c:\Users\AUSU\Documents\trae_projects\cc\tmp\petstore_er.png'

doc = Document(SRC)

# ============================================================
# 辅助：在元素前插入新元素
# ============================================================
def insert_before_element(new_element, ref_element):
    """在 ref_element 之前插入 new_element"""
    parent = ref_element.getparent()
    idx = list(parent).index(ref_element)
    parent.insert(idx, new_element)

def remove_elements_between(body, start_text, end_text):
    """删除正文中 start_text 段落与 end_text 段落之间的所有元素"""
    children = list(body)
    in_range = False
    to_remove = []
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag not in ('p', 'tbl'):
            continue

        # 获取段落文本
        text = ''
        if tag == 'p':
            text = ''.join(node.text or '' for node in child.iter(qn('w:t')))

        # 检测开始
        if not in_range and start_text in text:
            in_range = True
            to_remove.append(child)
            continue

        # 检测结束
        if in_range:
            if end_text in text:
                in_range = False
                break
            to_remove.append(child)

    for child in to_remove:
        body.remove(child)
    print(f'已删除 {len(to_remove)} 个元素（从"{start_text}"到"{end_text}"）')

# ============================================================
# 辅助：创建格式化段落
# ============================================================
def make_para(doc, text, style='Normal', bold=False, font_size=None, alignment=None, space_after=None):
    """创建段落"""
    p = doc.add_paragraph(text, style=style)
    if bold or font_size:
        for run in p.runs:
            if bold:
                run.bold = True
            if font_size:
                run.font.size = Pt(font_size)
    # 对新段落设置对齐
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def make_heading(doc, text, level=1):
    """创建标题"""
    h = doc.add_heading(text, level=level)
    return h

def set_cell_text(cell, text, bold=False, font_size=9, alignment=None):
    """设置单元格文本和格式"""
    # 清除现有内容
    for p in cell.paragraphs:
        for r in p.runs:
            r.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(font_size)
    run.font.name = 'Microsoft YaHei'
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment

def format_table(table):
    """统一表格样式"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 设置表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}></w:tblPr>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="999999"/>'
        '</w:tblBorders>'
    )
    # 移除旧边框
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblPr.append(borders)

def set_header_row(table, texts):
    """设置表头行（深蓝底白字）"""
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        set_cell_text(cell, text, bold=True, font_size=9)
        # 设置单元格背景色
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="2F5496" w:val="clear"/>'
        )
        cell._tc.get_or_add_tcPr().append(shading)
        # 字体白色
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# ============================================================
# 步骤1：找到"第二部分"段落，保存引用
# ============================================================
body = doc.element.body
children = list(body)
part2_para = None
for child in children:
    tag = child.tag.split('}')[-1]
    if tag == 'p':
        text = ''.join(node.text or '' for node in child.iter(qn('w:t')))
        if '第二部分' in text:
            part2_para = child
            break

if part2_para is None:
    print('错误：未找到"第二部分"标记段落')
    exit(1)

print('找到"第二部分"段落。')

# ============================================================
# 步骤2：删除第一部分所有内容（从"第一部分"到"第二部分"之前）
# ============================================================
children = list(body)
in_part1 = False
to_remove = []
for child in children:
    tag = child.tag.split('}')[-1]
    text = ''
    if tag == 'p':
        text = ''.join(node.text or '' for node in child.iter(qn('w:t')))

    if not in_part1 and '第一部分' in text and 'PetStore' in text:
        in_part1 = True
        to_remove.append(child)
        continue
    if in_part1:
        if child is part2_para:
            break
        to_remove.append(child)

for child in to_remove:
    body.remove(child)
print(f'已删除第一部分 {len(to_remove)} 个元素。')

# ============================================================
# 步骤3：在"第二部分"之前插入新内容
# 由于 add_paragraph 只能加在末尾，我们需要创建元素后插入
# ============================================================

# 先收集所有新建元素，再逐个插入到 part2_para 之前
new_elements = []

def add_new_para(text, style='Normal', bold=False, font_size=10.5, font_name='Microsoft YaHei',
                 alignment=None, space_after=6, space_before=0, first_line_indent=None):
    """创建新段落并返回其XML元素"""
    p = doc.add_paragraph()
    # 从文档末尾移除刚加的元素
    body.remove(p._element)
    new_elements.append(p._element)

    # 设置样式
    if style != 'Normal':
        p.style = doc.styles[style]
    run = p.add_run(text)
    run.font.size = Pt(font_size)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if bold:
        run.bold = True
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_new_heading(text, level=1):
    """创建标题"""
    p = doc.add_heading(text, level=level)
    body.remove(p._element)
    new_elements.append(p._element)
    return p

def add_new_table(rows, cols):
    """创建表格"""
    table = doc.add_table(rows=rows, cols=cols)
    body.remove(table._element)
    new_elements.append(table._element)
    format_table(table)
    return table

def add_new_image(img_path, width_inches=5.5):
    """插入图片"""
    p = doc.add_paragraph()
    body.remove(p._element)
    new_elements.append(p._element)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))
    p.paragraph_format.space_after = Pt(8)
    return p

def add_empty_para():
    """空行"""
    p = doc.add_paragraph()
    body.remove(p._element)
    new_elements.append(p._element)
    p.paragraph_format.space_after = Pt(2)
    return p


# ===== 构建新内容 =====

# 标题
add_new_heading('第一部分　PetStore数据库的设计', level=0)

# ---- 一、概念结构设计 ----
add_new_heading('一、概念结构设计（ER图）', level=1)

add_new_para('根据宠物商店电子商务系统的业务逻辑，对需求进行分析后，抽取出以下实体及其属性，'
             '并确定实体之间的联系，最终绘制 E-R 图。')

add_new_heading('1.1 实体分析', level=2)

add_new_para('系统包含 6 个实体：用户、商品分类、商品、购物车、订单、订单明细。各实体属性及主键如下表所示。')

# 表1-1: 实体及属性汇总表
t = add_new_table(7, 3)
set_header_row(t, ['实体', '包含属性', '主键'])
data_entity = [
    ['用户',     '用户号、用户名、密码、性别、邮箱、电话',                     '用户号'],
    ['商品分类', '分类编号、分类名称',                                        '分类编号'],
    ['商品',     '商品编号、商品名、商品介绍、市场价格、当前价格、数量、分类编号', '商品编号'],
    ['购物车',   '购物车编号、用户号、商品编号、数量',                          '购物车编号'],
    ['订单',     '订单号、用户号、订单日期、订单总价、是否已处理',               '订单号'],
    ['订单明细', '明细编号、订单号、商品编号、商品名、单价、数量',               '明细编号'],
]
for i, row_data in enumerate(data_entity):
    for j, val in enumerate(row_data):
        set_cell_text(t.rows[i+1].cells[j], val, font_size=9)

# 表标题
cap1 = add_new_para('表1-1　实体及属性汇总表', font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=10, space_before=4)

add_empty_para()

# 各实体详细说明
add_new_heading('1.2 实体属性详细说明', level=2)
entity_details = [
    '用户：用户号（唯一标识）、用户名、密码、性别、邮箱、电话——来源于用户注册功能。',
    '商品分类：分类编号（唯一标识）、分类名称——管理员使用，用于组织商品层级结构。',
    '商品：商品编号（唯一标识）、商品名、商品介绍、市场价格、当前价格、数量、分类编号（外键，关联商品分类）——管理员录入的商品信息。',
    '购物车：购物车编号（唯一标识）、用户号（外键，关联用户）、商品编号（外键，关联商品）、数量——记录用户临时选购的宠物商品。',
    '订单：订单号（唯一标识）、用户号（外键，关联用户）、订单日期、订单总价、是否已处理——用户完成选购后提交的预订记录。',
    '订单明细：明细编号（唯一标识）、订单号（外键，关联订单）、商品编号（外键，关联商品）、商品名、单价、数量——每张订单中每个商品的购买详情。',
]
for d in entity_details:
    add_new_para(d, font_size=10.5, first_line_indent=0.74)

add_empty_para()

# 联系分析
add_new_heading('1.3 联系分析', level=2)
add_new_para('实体之间存在 6 个一对多（1 : N）联系，汇总如下表。')

# 表1-2: 联系汇总表
t2 = add_new_table(7, 6)
set_header_row(t2, ['序号', '实体1', '实体2', '联系类型', '联系名称', '说明'])
rel_data = [
    ['1', '商品分类', '商品',   '1 : N', '包含',   '一个分类下可包含多个商品，一个商品只属于一个分类'],
    ['2', '用户',   '购物车',  '1 : N', '添加',   '一个用户可添加多条购物车记录，每条记录属于一个用户'],
    ['3', '商品',   '购物车',  '1 : N', '被加入', '一个商品可被多个用户加入购物车，每条记录对应一个商品'],
    ['4', '用户',   '订单',   '1 : N', '下达',   '一个用户可下达多张订单，每张订单只属于一个用户'],
    ['5', '订单',   '订单明细', '1 : N', '拥有',   '一张订单包含多条明细，每条明细只属于一张订单'],
    ['6', '商品',   '订单明细', '1 : N', '对应',   '一个商品可出现在多条订单明细中，每条明细对应一个商品'],
]
for i, row_data in enumerate(rel_data):
    for j, val in enumerate(row_data):
        set_cell_text(t2.rows[i+1].cells[j], val, font_size=9)

cap2 = add_new_para('表1-2　实体间联系汇总表', font_size=9, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    space_after=10, space_before=4)

add_empty_para()

# ER 图
add_new_heading('1.4 E-R 图', level=2)
add_new_para('PetStore 数据库的整体 E-R 结构如下图所示：', font_size=10.5)

add_new_image(ER_IMG, width_inches=5.8)

cap_img = add_new_para('图1-1　PetStore 数据库 E-R 图', font_size=9,
                       alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

add_empty_para()

# ER 图文字描述
add_new_para('系统核心围绕用户展开：管理员维护商品分类，每个分类下包含多种商品。'
             '用户注册后可按分类浏览商品，将心仪商品加入购物车（一条购物车记录关联一个用户和一个商品）。'
             '用户确认选购后提交订单，每张订单包含多条订单明细，每条明细记录所购商品、单价和数量，'
             '最终汇总生成订单总价。',
             font_size=10.5, first_line_indent=0.74)

# ---- 二、逻辑结构设计 ----
add_new_heading('二、逻辑结构设计（关系模型）', level=1)

add_new_para('将 E-R 图转换为关系模型：每个实体转换为一个关系（表），实体的属性即为表的字段，'
             '实体的主键即为表的主键。实体间的 1 : N 联系通过在 N 端表中添加外键实现。'
             '最终得到 6 张表，各表的逻辑结构如下。', font_size=10.5, first_line_indent=0.74)

add_empty_para()

# ===== 6 张表的结构定义 =====
tables_def = [
    ('表2-1　用户表（user）', [
        ['用户号',  'CHAR(10)',    'NOT NULL', 'PRIMARY KEY', '用户唯一标识'],
        ['用户名',  'VARCHAR(20)', 'NOT NULL', '—',          '用户昵称'],
        ['密码',    'VARCHAR(20)', 'NOT NULL', '—',          '登录密码'],
        ['性别',    'CHAR(2)',     'NULL',     '—',          '男 / 女'],
        ['邮箱',    'VARCHAR(50)', 'NULL',     '—',          '电子邮箱地址'],
        ['电话',    'CHAR(11)',    'NULL',     '—',          '联系电话'],
    ]),
    ('表2-2　商品分类表（category）', [
        ['分类编号', 'CHAR(10)',    'NOT NULL', 'PRIMARY KEY', '分类唯一标识'],
        ['分类名称', 'VARCHAR(20)', 'NOT NULL', '—',          '如"狗""猫""鱼""鸟"'],
    ]),
    ('表2-3　商品表（product）', [
        ['商品编号', 'CHAR(10)',       'NOT NULL', 'PRIMARY KEY',              '商品唯一标识'],
        ['商品名',   'VARCHAR(50)',    'NOT NULL', '—',                       '商品名称'],
        ['商品介绍', 'TEXT',           'NULL',     '—',                       '商品详细描述'],
        ['市场价格', 'DECIMAL(10,2)',  'NULL',     '—',                       '市场参考价'],
        ['当前价格', 'DECIMAL(10,2)',  'NOT NULL', '—',                       '本站实际售价'],
        ['数量',     'INT',            'NOT NULL', '—',                       '库存数量'],
        ['分类编号', 'CHAR(10)',       'NOT NULL', 'FOREIGN KEY',              '外键，关联 category(分类编号)'],
    ]),
    ('表2-4　购物车表（cart）', [
        ['购物车编号', 'INT',        'NOT NULL', 'PRIMARY KEY',              '自增，购物车记录唯一标识'],
        ['用户号',    'CHAR(10)',   'NOT NULL', 'FOREIGN KEY',              '外键，关联 user(用户号)'],
        ['商品编号',  'CHAR(10)',   'NOT NULL', 'FOREIGN KEY',              '外键，关联 product(商品编号)'],
        ['数量',      'INT',        'NOT NULL', '—',                       '选购数量，默认为 1'],
    ]),
    ('表2-5　订单表（order）', [
        ['订单号',    'CHAR(14)',      'NOT NULL', 'PRIMARY KEY',              '如日期+序号，唯一标识'],
        ['用户号',    'CHAR(10)',      'NOT NULL', 'FOREIGN KEY',              '外键，关联 user(用户号)'],
        ['订单日期',  'DATETIME',      'NOT NULL', '—',                       '下单时间'],
        ['订单总价',  'DECIMAL(10,2)', 'NOT NULL', '—',                       '订单金额合计'],
        ['是否已处理','TINYINT(1)',    'NOT NULL', 'DEFAULT 0',               '0=未处理，1=已处理'],
    ]),
    ('表2-6　订单明细表（order_detail）', [
        ['明细编号', 'INT',           'NOT NULL', 'PRIMARY KEY',              '自增，明细唯一标识'],
        ['订单号',   'CHAR(14)',      'NOT NULL', 'FOREIGN KEY',              '外键，关联 order(订单号)'],
        ['商品编号', 'CHAR(10)',      'NOT NULL', 'FOREIGN KEY',              '外键，关联 product(商品编号)'],
        ['商品名',   'VARCHAR(50)',   'NOT NULL', '—',                       '冗余存储，便于查询展示'],
        ['单价',     'DECIMAL(10,2)', 'NOT NULL', '—',                       '购买时的实际单价'],
        ['数量',     'INT',           'NOT NULL', '—',                       '购买数量'],
    ]),
]

for title, rows in tables_def:
    cap = add_new_para(title, font_size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       space_after=4, space_before=6)
    t = add_new_table(len(rows) + 1, 5)
    set_header_row(t, ['字段名', '数据类型', '允许为空', '约束', '说明'])

    # 设置列宽
    for row_obj in t.rows:
        for ci, w in enumerate([Cm(2.2), Cm(2.5), Cm(1.8), Cm(2.5), Cm(4.5)]):
            row_obj.cells[ci].width = w

    for i, row_data in enumerate(rows):
        for j, val in enumerate(row_data):
            set_cell_text(t.rows[i+1].cells[j], val, font_size=9)
    add_empty_para()

# ===== 关系模式总结 =====
add_new_heading('2.1 关系模式总结', level=2)
add_new_para('上述 6 张表的关系模式可简洁表示为（带下划线者为主键）：', font_size=10.5)

schemas = [
    '用户（用户号，用户名，密码，性别，邮箱，电话）',
    '商品分类（分类编号，分类名称）',
    '商品（商品编号，商品名，商品介绍，市场价格，当前价格，数量，分类编号）',
    '购物车（购物车编号，用户号，商品编号，数量）',
    '订单（订单号，用户号，订单日期，订单总价，是否已处理）',
    '订单明细（明细编号，订单号，商品编号，商品名，单价，数量）',
]
for s in schemas:
    add_new_para(s, font_size=10.5, first_line_indent=0.74, space_after=3)


# ============================================================
# 步骤4：将所有新元素插入到 part2_para 之前
# ============================================================
for elem in new_elements:
    insert_before_element(elem, part2_para)

print(f'已插入 {len(new_elements)} 个新元素。')

# ============================================================
# 步骤5：保存
# ============================================================
doc.save(DST)
print(f'文档已保存: {DST}')
print('完成！')
