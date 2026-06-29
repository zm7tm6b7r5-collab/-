r"""
品牌管理 PPT 图片 OCR → 结构化 Word 文档
批量处理 85 张PPT拍照图片
"""
import subprocess
import os
import re
import sys
import json
from pathlib import Path
from PIL import Image, ImageFilter, ImageEnhance
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Force UTF-8
sys.stdout.reconfigure(encoding='utf-8', errors='replace')

# ===== 配置 =====
IMG_DIR = Path(r"C:\Users\AUSU\Desktop\品牌")
OUTPUT_PATH = Path(r"C:\Users\AUSU\Desktop\品牌管理课程笔记.docx")
CACHE_PATH = Path(r"C:\Users\AUSU\Documents\trae_projects\cc\tmp\ocr_cache.json")
TESSERACT = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
TESSDATA = os.path.expandvars(r"%USERPROFILE%\tesseract-tessdata")


def load_cache() -> dict:
    """加载 OCR 缓存"""
    if CACHE_PATH.exists():
        try:
            with open(CACHE_PATH, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {}


def save_cache(cache: dict):
    """保存 OCR 缓存"""
    CACHE_PATH.parent.mkdir(parents=True, exist_ok=True)
    with open(CACHE_PATH, "w", encoding="utf-8") as f:
        json.dump(cache, f, ensure_ascii=False)


# ===== OCR 引擎 =====
def ocr_image(img_path: Path, cache: dict) -> str:
    """对图片进行 OCR，优先使用缓存"""
    cache_key = img_path.name
    if cache_key in cache:
        return cache[cache_key]

    try:
        img = Image.open(img_path)
        gray = img.convert("L")
        enhancer = ImageEnhance.Contrast(gray)
        enhanced = enhancer.enhance(2.5)
        enhanced = enhanced.filter(ImageFilter.SHARPEN)
        w, h = img.size
        if w < 1500:
            scale = max(1, 2000 // min(w, h))
            enhanced = enhanced.resize((w * scale, h * scale), Image.LANCZOS)
        tmp = IMG_DIR / f"_ocr_tmp_{img_path.stem}.png"
        enhanced.save(tmp)
        env = os.environ.copy()
        env["TESSDATA_PREFIX"] = TESSDATA
        r = subprocess.run(
            [TESSERACT, str(tmp), "stdout", "-l", "chi_sim+eng", "--psm", "6"],
            capture_output=True, text=True, env=env, timeout=60, encoding="utf-8"
        )
        tmp.unlink()
        text = r.stdout.strip()
        cache[cache_key] = text
        return text
    except Exception as e:
        print(f"  [ERROR] {img_path.name}: {e}")
        return ""


# ===== 文本清理 =====
def meaningful_ratio(line: str) -> float:
    """计算一行中有意义字符的比例（中文+英文+数字）"""
    if not line:
        return 0
    meaningful = sum(1 for c in line
                     if '一' <= c <= '鿿'  # 中文
                     or 'a' <= c.lower() <= 'z'     # 英文
                     or '0' <= c <= '9')             # 数字
    return meaningful / len(line)


def chinese_count(line: str) -> int:
    """计算中文字符数"""
    return sum(1 for c in line if '一' <= c <= '鿿')


def is_noise_line(line: str) -> bool:
    """判断是否为 OCR 噪声行"""
    stripped = line.strip()
    if not stripped:
        return False  # 空行不是噪声，是有意义的分隔

    # 极短线噪声（1-2字符且无中文）
    if len(stripped) <= 2 and chinese_count(stripped) == 0:
        return True

    # 纯符号行
    ratio = meaningful_ratio(stripped)
    if ratio < 0.15:
        return True

    # 纯数字/标点噪声
    if chinese_count(stripped) == 0 and len(stripped) < 5 and ratio < 0.5:
        return True

    return False


def clean_text(text: str) -> list:
    """清理 OCR 文本，返回有意义的行列表（保留空行作分隔）"""
    lines = text.split("\n")
    result = []
    consecutive_empty = 0

    for line in lines:
        stripped = line.strip()

        if is_noise_line(line):
            continue

        if not stripped:
            consecutive_empty += 1
            if consecutive_empty <= 1:  # 最多保留一个空行
                result.append("")
            continue

        consecutive_empty = 0
        result.append(stripped)

    # 去除首尾空行
    while result and not result[0]:
        result.pop(0)
    while result and not result[-1]:
        result.pop()

    return result


# ===== 标题和内容检测 =====
def looks_like_heading(line: str) -> bool:
    """
    严格判断是否为标题行。
    条件严格：必须有足够的中文内容 + 明确的标题特征
    """
    cn = chinese_count(line)
    total = len(line)

    # 没有中文：只有纯英文全大写短行才算标题
    if cn == 0:
        # 全大写英文，5-50字符
        if line == line.upper() and 5 <= total <= 50 and line[0].isalpha():
            return True
        return False

    # 必须有至少3个中文字符
    if cn < 3:
        return False

    # 有意义字符比例必须 > 50%（排除符号噪声）
    if meaningful_ratio(line) < 0.5:
        return False

    # 1. "一、"/"1."/"1)" 序号 + 中文
    if re.match(r'^[一二三四五六七八九十\d]+[、.)．]\s*[一-鿿]', line):
        return True

    # 2. "第X章/X讲" 模式
    if re.match(r'^第[一二三四五六七八九十\d]+[章节讲课]', line):
        return True

    # 3. 短中文行（5-15字），纯中文无符号噪声
    if 5 <= cn <= 15 and cn / total >= 0.7:
        return True

    # 4. "XXX：XXX" 格式，冒号前 <= 10字
    if '：' in line:
        parts = line.split('：')
        if len(parts[0]) <= 10 and chinese_count(parts[0]) >= 1:
            return True

    return False


def looks_like_english_block(lines: list) -> bool:
    """判断行列表是否主要是英文块"""
    if not lines:
        return False
    total_ratio = sum(meaningful_ratio(l) for l in lines) / len(lines)
    chinese_total = sum(chinese_count(l) for l in lines)
    # 英文为主：有意义字符占比高且中文极少
    return total_ratio > 0.6 and chinese_total <= 2


# ===== Word 文档生成 =====
def create_docx(all_texts: list, doc_path: Path):
    """生成结构化 Word 文档"""
    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)

    # 默认样式
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.5

    # === 封面 ===
    title = doc.add_heading("品牌管理课程笔记", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"共 {len(all_texts)} 页 PPT | OCR 自动识别整理 | 2026年6月26日")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_page_break()

    # === 正文 ===
    slide_num = 0
    for orig_idx, raw_text in enumerate(all_texts):
        lines = clean_text(raw_text)
        if not lines:
            continue

        slide_num += 1

        # 页间分隔
        if slide_num > 1:
            doc.add_paragraph()

        # 页码标记
        page_marker = doc.add_paragraph()
        page_marker.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = page_marker.add_run(f"-- Slide {slide_num} --")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(180, 180, 180)
        run.font.italic = True

        # 逐行处理
        i = 0
        while i < len(lines):
            line = lines[i]
            if not line:
                doc.add_paragraph()  # 空行 = 段落分隔
                i += 1
                continue

            # 收集连续的英文行
            if meaningful_ratio(line) > 0.7 and chinese_count(line) == 0:
                english_batch = []
                while i < len(lines) and lines[i] and meaningful_ratio(lines[i]) > 0.7 and chinese_count(lines[i]) == 0:
                    english_batch.append(lines[i])
                    i += 1
                # 输出英文块
                for el in english_batch:
                    p = doc.add_paragraph()
                    run = p.add_run(el)
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(90, 90, 90)
                continue

            # 判断是否为标题
            if looks_like_heading(line):
                doc.add_heading(line, level=2)
                i += 1
                continue

            # 普通正文
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(11)
            i += 1

    # === 尾注 ===
    doc.add_page_break()
    doc.add_heading("关于本笔记", level=1)
    note = doc.add_paragraph()
    run_note = note.add_run("整理说明：\n").bold = True
    note.add_run(
        "1. 本笔记由 Tesseract OCR 自动识别 PPT 拍照图片生成，可能存在识别错误。\n"
        "2. 建议对照原始图片（位于「桌面/品牌」文件夹）核对和修正。\n"
        "3. 英文部分为 OCR 识别结果，未做翻译校对，以灰色小字标出。\n"
        "4. 整理日期：2026年6月26日\n"
    )

    # 保存
    doc.save(str(doc_path))
    print(f"\n  [OK] Document saved: {doc_path}")


# ===== 主流程 =====
def main():
    print("=" * 60)
    print("  Brand Management PPT -> Word OCR Converter")
    print("=" * 60)

    # 获取所有图片
    images = sorted(IMG_DIR.glob("*.jpg"))
    print(f"\n[Images] Found {len(images)} images")

    # 加载缓存
    cache = load_cache()
    print(f"[Cache] {len(cache)} cached entries loaded")

    # 批量 OCR
    all_texts = []
    for idx, img_path in enumerate(images, 1):
        cached = " (cached)" if img_path.name in cache else ""
        print(f"  [{idx:2d}/{len(images)}] {img_path.name}{cached} ...", end=" ", flush=True)
        text = ocr_image(img_path, cache)
        chars = len(text) if text else 0
        print(f"{chars:4d} chars")
        all_texts.append(text)
        # 每10张保存一次缓存
        if idx % 10 == 0:
            save_cache(cache)

    # 统计
    total_chars = sum(len(t) for t in all_texts)
    non_empty = sum(1 for t in all_texts if t.strip())
    print(f"\n[Stats] {non_empty}/{len(images)} pages with text, ~{total_chars:,} chars")

    # 生成 Word
    print(f"\n[Build] Generating Word document...")
    create_docx(all_texts, OUTPUT_PATH)

    # 清理临时文件
    print(f"\n[Clean] Removing temp files...")
    for tmp in IMG_DIR.glob("_ocr_tmp_*.png"):
        try:
            tmp.unlink()
        except Exception:
            pass
    print(f"  Done!")

    print(f"\n{'=' * 60}")
    print(f"  [DONE] All complete!")
    print(f"  Output: {OUTPUT_PATH}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
